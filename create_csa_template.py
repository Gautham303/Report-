#!/usr/bin/env python3
"""
CSA Annual Report Word Template Generator

This script creates a Word document template that replicates the design elements,
layout, and formatting from the CSA Annual Report PDF.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement, ns
from docx.oxml.shared import qn
import io
from PIL import Image, ImageDraw
import os

class CSATemplateGenerator:
    def __init__(self):
        self.doc = Document()
        self.setup_document()
        self.setup_styles()
        
        # Color scheme
        self.colors = {
            'dark_blue': RGBColor(46, 58, 135),      # #2E3A87
            'light_blue': RGBColor(0, 123, 191),     # #007BBF
            'teal': RGBColor(0, 150, 136),           # #009688
            'gray': RGBColor(128, 128, 128),         # #808080
            'light_gray': RGBColor(240, 240, 240),   # #F0F0F0
            'black': RGBColor(0, 0, 0),              # #000000
            'white': RGBColor(255, 255, 255)         # #FFFFFF
        }
    
    def setup_document(self):
        """Configure document settings"""
        # Set page size to A4
        section = self.doc.sections[0]
        section.page_height = Inches(11.7)  # A4 height
        section.page_width = Inches(8.3)    # A4 width
        
        # Set margins
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    def setup_styles(self):
        """Create custom styles for the document"""
        styles = self.doc.styles
        
        # Cover page title style
        try:
            cover_title = styles.add_style('CoverTitle', WD_STYLE_TYPE.PARAGRAPH)
            cover_title.font.name = 'Arial'
            cover_title.font.size = Pt(48)
            cover_title.font.bold = True
            cover_title.font.color.rgb = self.colors['white']
            cover_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cover_title.paragraph_format.space_after = Pt(12)
        except:
            pass  # Style might already exist
            
        # Cover page subtitle style
        try:
            cover_subtitle = styles.add_style('CoverSubtitle', WD_STYLE_TYPE.PARAGRAPH)
            cover_subtitle.font.name = 'Arial'
            cover_subtitle.font.size = Pt(36)
            cover_subtitle.font.bold = True
            cover_subtitle.font.color.rgb = self.colors['dark_blue']
            cover_subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cover_subtitle.paragraph_format.space_after = Pt(6)
        except:
            pass
            
        # Header style
        try:
            header_style = styles.add_style('HeaderStyle', WD_STYLE_TYPE.PARAGRAPH)
            header_style.font.name = 'Arial'
            header_style.font.size = Pt(12)
            header_style.font.bold = True
            header_style.font.color.rgb = self.colors['white']
            header_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        except:
            pass
            
        # Section title style
        try:
            section_title = styles.add_style('SectionTitle', WD_STYLE_TYPE.PARAGRAPH)
            section_title.font.name = 'Arial'
            section_title.font.size = Pt(24)
            section_title.font.bold = True
            section_title.font.color.rgb = self.colors['dark_blue']
            section_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            section_title.paragraph_format.space_after = Pt(12)
        except:
            pass
            
        # TOC heading style
        try:
            toc_heading = styles.add_style('TOCHeading', WD_STYLE_TYPE.PARAGRAPH)
            toc_heading.font.name = 'Arial'
            toc_heading.font.size = Pt(36)
            toc_heading.font.bold = True
            toc_heading.font.color.rgb = self.colors['dark_blue']
            toc_heading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            toc_heading.paragraph_format.space_after = Pt(24)
        except:
            pass
            
        # TOC item style
        try:
            toc_item = styles.add_style('TOCItem', WD_STYLE_TYPE.PARAGRAPH)
            toc_item.font.name = 'Arial'
            toc_item.font.size = Pt(14)
            toc_item.font.color.rgb = self.colors['black']
            toc_item.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            toc_item.paragraph_format.space_after = Pt(8)
        except:
            pass
            
        # Body text style
        try:
            body_text = styles.add_style('BodyText', WD_STYLE_TYPE.PARAGRAPH)
            body_text.font.name = 'Calibri'
            body_text.font.size = Pt(11)
            body_text.font.color.rgb = self.colors['black']
            body_text.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            body_text.paragraph_format.space_after = Pt(12)
            body_text.paragraph_format.line_spacing = 1.15
        except:
            pass

    def create_placeholder_image(self, width_px=400, height_px=300, text="PLACEHOLDER IMAGE"):
        """Create a placeholder image"""
        img = Image.new('RGB', (width_px, height_px), color=(200, 200, 200))
        draw = ImageDraw.Draw(img)
        
        # Add placeholder text
        bbox = draw.textbbox((0, 0), text)
        text_width = bbox[2] - bbox[0]
        text_height = bbox[3] - bbox[1]
        
        x = (width_px - text_width) // 2
        y = (height_px - text_height) // 2
        draw.text((x, y), text, fill=(100, 100, 100))
        
        # Save to memory
        img_io = io.BytesIO()
        img.save(img_io, format='PNG')
        img_io.seek(0)
        return img_io

    def create_cover_page(self):
        """Create the cover page with design elements"""
        # Add cover page background note
        p = self.doc.add_paragraph()
        p.add_run("DESIGN NOTE: Apply dark blue gradient background with geometric diamond overlay manually in Word").italic = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add spacing
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        
        # Main title
        title = self.doc.add_paragraph("COMPUTER SCIENCE ASSOCIATION", style='CoverTitle')
        
        # Year with blue accent
        year_para = self.doc.add_paragraph()
        year_run = year_para.add_run("2023-24")
        year_run.font.name = 'Arial'
        year_run.font.size = Pt(32)
        year_run.font.bold = True
        year_run.font.color.rgb = self.colors['dark_blue']
        year_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add spacing
        self.doc.add_paragraph()
        
        # Annual Report title
        report_title = self.doc.add_paragraph()
        report_run = report_title.add_run("ANNUAL REPORT")
        report_run.font.name = 'Arial'
        report_run.font.size = Pt(28)
        report_run.font.bold = True
        report_run.font.color.rgb = self.colors['black']
        report_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add spacing
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        
        # Placeholder for diamond-shaped image frame
        diamond_note = self.doc.add_paragraph()
        diamond_note.add_run("DESIGN NOTE: Insert diamond-shaped image frame with college building photo here").italic = True
        diamond_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add placeholder image
        img_stream = self.create_placeholder_image(300, 300, "COLLEGE BUILDING")
        self.doc.add_picture(img_stream, width=Inches(3))
        last_paragraph = self.doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add spacing and design elements note
        self.doc.add_paragraph()
        design_note = self.doc.add_paragraph()
        design_note.add_run("DESIGN NOTE: Add modern geometric design elements in blue/teal colors manually").italic = True
        design_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Page break
        self.doc.add_page_break()

    def setup_headers_footers(self):
        """Setup headers and footers for interior pages"""
        section = self.doc.sections[0]
        header = section.header
        footer = section.footer
        
        # Clear existing header/footer
        header.paragraphs[0].clear()
        
        # Create header with CS Association and page number
        header_para = header.paragraphs[0]
        
        # Left side - CS Association in blue box
        left_run = header_para.add_run("CS ASSOCIATION")
        left_run.font.name = 'Arial'
        left_run.font.size = Pt(10)
        left_run.font.bold = True
        left_run.font.color.rgb = self.colors['white']
        
        # Add tab to right align page number
        tab_run = header_para.add_run("\t\t\t\t\t\t")
        
        # Right side - Page number in black box
        page_run = header_para.add_run("Page ")
        page_run.font.name = 'Arial'
        page_run.font.size = Pt(10)
        page_run.font.bold = True
        page_run.font.color.rgb = self.colors['black']
        
        # Add instruction for manual formatting
        note_para = header.add_paragraph()
        note_para.add_run("DESIGN NOTE: Format 'CS ASSOCIATION' with blue background box, page number with black background box").italic = True

    def create_table_of_contents(self):
        """Create table of contents page"""
        # TOC Title with blue accent line
        toc_title = self.doc.add_paragraph("Table Of Content", style='TOCHeading')
        
        # Add blue accent line note
        accent_note = self.doc.add_paragraph()
        accent_note.add_run("DESIGN NOTE: Add blue horizontal accent line under title manually").italic = True
        
        self.doc.add_paragraph()
        
        # Create table for two-column layout
        table = self.doc.add_table(rows=0, cols=2)
        table.autofit = False
        table.columns[0].width = Inches(3)
        table.columns[1].width = Inches(3)
        
        # TOC items with circular numbered badges
        toc_items = [
            ("03", "About CS Association"),
            ("05", "President's Message"),
            ("06", "Vision & Mission"),
            ("08", "Department Overview"),
            ("10", "Faculty Members"),
            ("12", "Student Activities"),
            ("15", "Technical Events"),
            ("18", "Cultural Programs"),
            ("20", "Achievements"),
            ("22", "Industry Partnerships"),
            ("25", "Alumni Network"),
            ("28", "Future Plans"),
            ("30", "Acknowledgments"),
            ("32", "Contact Information")
        ]
        
        # Add items to table (2 items per row)
        for i in range(0, len(toc_items), 2):
            row = table.add_row()
            
            # Left column
            left_cell = row.cells[0]
            left_para = left_cell.paragraphs[0]
            left_para.add_run(f"⚫ {toc_items[i][0]}  {toc_items[i][1]}")
            left_para.style = 'TOCItem'
            
            # Right column (if exists)
            if i + 1 < len(toc_items):
                right_cell = row.cells[1]
                right_para = right_cell.paragraphs[0]
                right_para.add_run(f"⚫ {toc_items[i+1][0]}  {toc_items[i+1][1]}")
                right_para.style = 'TOCItem'
        
        # Background image integration note
        self.doc.add_paragraph()
        bg_note = self.doc.add_paragraph()
        bg_note.add_run("DESIGN NOTE: Add subtle background image integration manually").italic = True
        bg_note.add_run("\nDESIGN NOTE: Replace ⚫ with circular blue numbered badges manually").italic = True
        
        # Page break
        self.doc.add_page_break()

    def create_content_pages(self):
        """Create sample content pages"""
        
        # About CS Association section
        section_title = self.doc.add_paragraph("About CS Association", style='SectionTitle')
        
        # Blue accent line note
        accent_note = self.doc.add_paragraph()
        accent_note.add_run("DESIGN NOTE: Add blue vertical accent line on left margin manually").italic = True
        
        self.doc.add_paragraph()
        
        # Sample content with professional formatting
        content_paragraphs = [
            "The Computer Science Association (CSA) is a vibrant student organization dedicated to fostering academic excellence, professional development, and community engagement within the field of computer science. Established with the vision of creating a collaborative learning environment, CSA serves as a bridge between theoretical knowledge and practical application.",
            
            "Our association brings together students, faculty, and industry professionals to create meaningful connections and opportunities. Through various initiatives, workshops, seminars, and events, we strive to enhance the educational experience and prepare our members for successful careers in technology.",
            
            "CSA is committed to promoting innovation, research, and entrepreneurship among computer science students. We believe in the power of collaboration and the importance of staying current with emerging technologies and industry trends.",
            
            "Throughout the academic year, our association organizes numerous activities including technical workshops, coding competitions, guest lectures from industry experts, project showcases, and networking events. These activities are designed to complement classroom learning and provide practical exposure to real-world challenges."
        ]
        
        for content in content_paragraphs:
            para = self.doc.add_paragraph(content, style='BodyText')
        
        # Image integration placeholder
        self.doc.add_paragraph()
        img_note = self.doc.add_paragraph()
        img_note.add_run("DESIGN NOTE: Insert relevant images within content sections with proper alignment").italic = True
        
        # Add placeholder image
        img_stream = self.create_placeholder_image(400, 250, "CONTENT IMAGE")
        self.doc.add_picture(img_stream, width=Inches(4))
        last_paragraph = self.doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Continue with more content
        self.doc.add_paragraph()
        
        more_content = [
            "Our mission extends beyond academic pursuits to include community service and social responsibility. We actively participate in outreach programs, mentor junior students, and contribute to open-source projects that benefit the broader technology community.",
            
            "The association also focuses on professional development through resume workshops, interview preparation sessions, and career guidance programs. We maintain strong relationships with alumni who provide valuable insights and mentorship opportunities."
        ]
        
        for content in more_content:
            para = self.doc.add_paragraph(content, style='BodyText')
        
        # Page break
        self.doc.add_page_break()

    def create_organizational_chart(self):
        """Create organizational chart template"""
        
        # Page title
        org_title = self.doc.add_paragraph("Organizational Structure", style='SectionTitle')
        
        self.doc.add_paragraph()
        
        # Hierarchical structure note
        hierarchy_note = self.doc.add_paragraph()
        hierarchy_note.add_run("DESIGN NOTE: Create hierarchical structure with connecting lines manually").italic = True
        
        self.doc.add_paragraph()
        
        # Executive Board
        exec_heading = self.doc.add_paragraph()
        exec_run = exec_heading.add_run("Executive Board")
        exec_run.font.name = 'Arial'
        exec_run.font.size = Pt(16)
        exec_run.font.bold = True
        exec_run.font.color.rgb = self.colors['dark_blue']
        
        # President
        president_table = self.doc.add_table(rows=1, cols=3)
        president_table.autofit = False
        
        # Center the president
        president_row = president_table.rows[0]
        president_row.cells[1].text = "⭕ PRESIDENT\n[Name Placeholder]\nFinal Year CSE"
        president_row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_paragraph()
        
        # Vice Presidents and Secretary
        vp_table = self.doc.add_table(rows=1, cols=3)
        vp_table.autofit = False
        
        positions = ["⭕ VICE PRESIDENT\n(Technical)\n[Name Placeholder]", 
                    "⭕ SECRETARY\n[Name Placeholder]\nThird Year CSE",
                    "⭕ VICE PRESIDENT\n(Operations)\n[Name Placeholder]"]
        
        for i, position in enumerate(positions):
            vp_table.rows[0].cells[i].text = position
            vp_table.rows[0].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_paragraph()
        
        # Department Heads
        dept_heading = self.doc.add_paragraph()
        dept_run = dept_heading.add_run("Department Heads")
        dept_run.font.name = 'Arial'
        dept_run.font.size = Pt(14)
        dept_run.font.bold = True
        dept_run.font.color.rgb = self.colors['dark_blue']
        
        # Department heads table
        dept_table = self.doc.add_table(rows=2, cols=3)
        dept_table.autofit = False
        
        dept_positions = [
            ["⭕ TECHNICAL HEAD\n[Name Placeholder]\nSecond Year CSE", 
             "⭕ EVENT HEAD\n[Name Placeholder]\nSecond Year CSE",
             "⭕ CULTURAL HEAD\n[Name Placeholder]\nSecond Year CSE"],
            ["⭕ MEDIA HEAD\n[Name Placeholder]\nFirst Year CSE",
             "⭕ FINANCE HEAD\n[Name Placeholder]\nSecond Year CSE", 
             "⭕ OUTREACH HEAD\n[Name Placeholder]\nFirst Year CSE"]
        ]
        
        for row_idx, row_positions in enumerate(dept_positions):
            for col_idx, position in enumerate(row_positions):
                dept_table.rows[row_idx].cells[col_idx].text = position
                dept_table.rows[row_idx].cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Design notes
        self.doc.add_paragraph()
        design_notes = self.doc.add_paragraph()
        design_notes.add_run("DESIGN NOTES:").bold = True
        design_notes.add_run("\n• Replace ⭕ with circular profile photo layouts")
        design_notes.add_run("\n• Add connecting lines between positions to show hierarchy")  
        design_notes.add_run("\n• Use consistent spacing and alignment")
        design_notes.add_run("\n• Apply professional formatting to role titles and names")
        design_notes.paragraph_format.space_after = Pt(12)
        
        # Page break
        self.doc.add_page_break()

    def create_additional_templates(self):
        """Create additional page templates"""
        
        # Events page template
        events_title = self.doc.add_paragraph("Technical Events", style='SectionTitle')
        
        # Events table
        events_table = self.doc.add_table(rows=4, cols=3)
        events_table.style = 'Table Grid'
        
        # Header row
        header_cells = events_table.rows[0].cells
        header_cells[0].text = "Event Name"
        header_cells[1].text = "Date"
        header_cells[2].text = "Description"
        
        # Sample events
        events_data = [
            ["Code Quest 2024", "March 15, 2024", "Annual coding competition with multiple rounds"],
            ["Tech Talk Series", "Monthly", "Guest lectures from industry professionals"],
            ["Project Showcase", "May 20, 2024", "Student project presentations and demonstrations"]
        ]
        
        for i, event_data in enumerate(events_data, 1):
            row_cells = events_table.rows[i].cells
            for j, data in enumerate(event_data):
                row_cells[j].text = data
        
        self.doc.add_paragraph()
        
        # Achievements section
        achieve_title = self.doc.add_paragraph("Key Achievements", style='SectionTitle')
        
        achievements = [
            "🏆 Won 1st place in Inter-College Programming Contest",
            "🏆 Successfully organized 5 technical workshops with 200+ participants",
            "🏆 Established partnerships with 3 major tech companies",
            "🏆 Launched mentorship program connecting 50+ students with alumni",
            "🏆 Contributed to 10+ open-source projects as a community"
        ]
        
        for achievement in achievements:
            bullet_para = self.doc.add_paragraph()
            bullet_run = bullet_para.add_run(achievement)
            bullet_run.font.name = 'Calibri'
            bullet_run.font.size = Pt(12)
            bullet_para.paragraph_format.left_indent = Inches(0.25)
            bullet_para.paragraph_format.space_after = Pt(6)
        
        self.doc.add_paragraph()
        
        # Contact information
        contact_title = self.doc.add_paragraph("Contact Information", style='SectionTitle')
        
        contact_info = self.doc.add_paragraph()
        contact_info.add_run("Computer Science Association\n")
        contact_info.add_run("Department of Computer Science & Engineering\n")  
        contact_info.add_run("[College Name]\n")
        contact_info.add_run("[Address Line 1]\n")
        contact_info.add_run("[Address Line 2]\n\n")
        contact_info.add_run("Email: csa@college.edu\n")
        contact_info.add_run("Phone: +91-XXXX-XXXXX\n")
        contact_info.add_run("Website: www.college.edu/csa\n")
        contact_info.add_run("Social Media: @CSACollege")
        
        contact_info.style = 'BodyText'

    def generate_template(self, output_path):
        """Generate the complete Word template"""
        print("Creating CSA Annual Report Template...")
        
        # Create all sections
        self.create_cover_page()
        self.create_table_of_contents()  
        self.create_content_pages()
        self.create_organizational_chart()
        self.create_additional_templates()
        
        # Setup headers and footers (applies to all pages after cover)
        self.setup_headers_footers()
        
        # Save the document
        self.doc.save(output_path)
        print(f"Template saved to: {output_path}")
        
        # Create documentation
        self.create_documentation()
        
    def create_documentation(self):
        """Create documentation for using the template"""
        doc_content = """
# CSA Annual Report Template Documentation

## Overview
This Word document template replicates the design elements, layout, and formatting from the CSA Annual Report PDF. The template includes all necessary sections and styling for creating a professional annual report.

## Template Sections

### 1. Cover Page
- **Design Elements**: Dark blue gradient background with geometric diamond overlay
- **Content**: Main title, year with blue accent, "Annual Report" subtitle
- **Image**: Diamond-shaped frame for college building photo
- **Manual Formatting Required**: 
  - Apply gradient background
  - Add geometric overlay elements
  - Format image in diamond shape

### 2. Table of Contents
- **Layout**: Two-column format with numbered badges
- **Manual Formatting Required**:
  - Replace bullet points with circular blue numbered badges
  - Add background image integration
  - Apply blue accent line under heading

### 3. Content Pages
- **Header**: CS Association in blue box (left), page number in black box (right)
- **Layout**: Professional paragraph formatting with blue accent lines
- **Manual Formatting Required**:
  - Format header boxes with appropriate colors
  - Add blue vertical accent line on left margin
  - Insert relevant images within content

### 4. Organizational Chart
- **Structure**: Hierarchical layout with connecting lines
- **Manual Formatting Required**:
  - Replace placeholder circles with actual profile photos
  - Add connecting lines between positions
  - Apply consistent spacing and professional formatting

## Color Scheme
- **Primary Blue**: #2E3A87 (Dark Blue)
- **Secondary Blue**: #007BBF (Light Blue)  
- **Accent**: #009688 (Teal)
- **Text**: #000000 (Black)
- **Background**: #FFFFFF (White)
- **Gray**: #808080 for accents

## Typography
- **Headers**: Arial, Bold
- **Body Text**: Calibri, 11pt
- **Titles**: Arial, various sizes as specified

## Usage Instructions

### Customizing Content
1. Replace placeholder text with actual content
2. Update [Name Placeholder] entries with real names and details
3. Insert actual photos in place of placeholder images
4. Modify dates, events, and achievements as needed

### Manual Formatting Tasks
1. **Cover Page**: Apply gradient background and geometric elements
2. **Headers**: Format CS Association and page number boxes
3. **TOC**: Replace bullets with blue numbered badges
4. **Accent Lines**: Add blue accent lines as indicated
5. **Org Chart**: Insert photos and connecting lines
6. **Images**: Replace placeholders with actual photos

### Adding New Sections
- Use the existing styles for consistency
- Follow the header/footer pattern for new pages
- Maintain the color scheme and typography
- Add blue accent lines for section headings

## Technical Notes
- Document size: A4 (8.3" x 11.7")
- Margins: 1 inch on all sides
- Line spacing: 1.15 for body text
- All placeholder elements are marked with "DESIGN NOTE" for easy identification

This template provides the foundation for creating a professional annual report matching the original PDF design.
"""
        
        with open('/home/runner/work/Report-/Report-/TEMPLATE_DOCUMENTATION.md', 'w') as f:
            f.write(doc_content)
        
        print("Documentation created: TEMPLATE_DOCUMENTATION.md")

def main():
    """Main function to generate the template"""
    generator = CSATemplateGenerator()
    output_path = '/home/runner/work/Report-/Report-/CSA_Annual_Report_Template.docx'
    generator.generate_template(output_path)
    
    print("\n" + "="*50)
    print("CSA Annual Report Template Generation Complete!")
    print("="*50)
    print(f"Template file: CSA_Annual_Report_Template.docx")
    print(f"Documentation: TEMPLATE_DOCUMENTATION.md")
    print("\nNext steps:")
    print("1. Open the template in Microsoft Word")
    print("2. Follow the 'DESIGN NOTE' instructions for manual formatting")
    print("3. Replace placeholder content with actual information")
    print("4. Add images, photos, and design elements as specified")
    print("5. Review and customize colors, fonts, and layout as needed")

if __name__ == "__main__":
    main()